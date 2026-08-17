"""
loaders_factory.py

Unified document loading for StudyMate AI.

Why this file was rewritten
----------------------------
The original implementation used PyPDFLoader / Docx2txtLoader directly.
Both flatten tables into plain prose text (columns get smashed together
with no separators), which silently destroys tabular data before it ever
reaches the retriever. For a study app, where source material routinely
contains comparison tables, rubrics, data tables, syllabi, etc., that's a
real accuracy bug, not a cosmetic one.

This version:
- Extracts prose text and tables as SEPARATE Document objects.
- Converts every extracted table into a Markdown table string, which an
  LLM can read far more reliably than collapsed, separator-less text, and
  which the splitter can recognize and never break apart.
- Adds unified metadata across every file type: `source`, `page`,
  `file_type`, `content_type` ("text" | "table").
"""

from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document


def _table_to_markdown(table_rows: list[list]) -> str:
    """Convert a raw row/column table (list of lists of cell text) into Markdown."""

    rows = [
        [("" if cell is None else str(cell).strip()) for cell in row]
        for row in table_rows
        if row
    ]

    if not rows:
        return ""

    header, *body = rows
    col_count = len(header)

    def pad(row):
        row = list(row) + [""] * (col_count - len(row))
        return row[:col_count]

    lines = [
        "| " + " | ".join(pad(header)) + " |",
        "| " + " | ".join(["---"] * col_count) + " |",
    ]

    for row in body:
        lines.append("| " + " | ".join(pad(row)) + " |")

    return "\n".join(lines)


class DocumentLoader:
    """
    Factory class responsible for loading different document types into
    a unified list of LangChain Document objects, with tables kept intact
    as separate, clearly-tagged Documents.
    """

    @staticmethod
    def load(file_path: str) -> list[Document]:

        extension = Path(file_path).suffix.lower()

        if extension == ".pdf":
            return DocumentLoader._load_pdf(file_path)

        if extension == ".txt":
            return DocumentLoader._load_txt(file_path)

        if extension == ".docx":
            return DocumentLoader._load_docx(file_path)

        raise ValueError(f"Unsupported file type: {extension}")

    # ------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------

    @staticmethod
    def _load_pdf(file_path: str) -> list[Document]:

        source = Path(file_path).name
        documents = []

        with pdfplumber.open(file_path) as pdf:

            for page_number, page in enumerate(pdf.pages, start=1):

                found_tables = page.find_tables()

                for table_index, found_table in enumerate(found_tables, start=1):

                    markdown_table = _table_to_markdown(found_table.extract())

                    if not markdown_table.strip():
                        continue

                    documents.append(
                        Document(
                            page_content=markdown_table,
                            metadata={
                                "source": source,
                                "page": page_number,
                                "file_type": "pdf",
                                "content_type": "table",
                                "table_index": table_index,
                            },
                        )
                    )

                # Exclude every table's bounding box from the prose text
                # extraction below. Without this, pdfplumber's page-level
                # extract_text() re-flattens the same cells back into the
                # text stream, so every table would appear twice: once as
                # a clean Markdown table and once as garbled inline text.
                table_bboxes = [t.bbox for t in found_tables]

                def outside_all_tables(obj, bboxes=table_bboxes):
                    for (tx0, ttop, tx1, tbottom) in bboxes:
                        if (
                            obj["x0"] >= tx0
                            and obj["x1"] <= tx1
                            and obj["top"] >= ttop
                            and obj["bottom"] <= tbottom
                        ):
                            return False
                    return True

                text_page = page.filter(outside_all_tables) if table_bboxes else page
                text = text_page.extract_text() or ""

                if text.strip():
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": source,
                                "page": page_number,
                                "file_type": "pdf",
                                "content_type": "text",
                            },
                        )
                    )

        return documents

    # ------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------

    @staticmethod
    def _load_docx(file_path: str) -> list[Document]:

        source = Path(file_path).name
        docx_file = DocxDocument(file_path)
        documents = []
        table_index = 0
        paragraph_lines: list[str] = []

        def flush_paragraphs():
            nonlocal paragraph_lines
            if paragraph_lines:
                text = "\n".join(paragraph_lines).strip()
                if text:
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": source,
                                "page": 0,
                                "file_type": "docx",
                                "content_type": "text",
                            },
                        )
                    )
                paragraph_lines = []

        # Walk the raw XML body so text and tables come out in roughly the
        # same order they appear in the document (python-docx normally
        # exposes .paragraphs and .tables as two separate, unordered lists).
        for child in docx_file.element.body.iterchildren():

            if isinstance(child, CT_P):
                paragraph = DocxParagraph(child, docx_file)
                if paragraph.text.strip():
                    paragraph_lines.append(paragraph.text)

            elif isinstance(child, CT_Tbl):
                flush_paragraphs()

                table = DocxTable(child, docx_file)
                table_index += 1

                rows = [[cell.text for cell in row.cells] for row in table.rows]
                markdown_table = _table_to_markdown(rows)

                if markdown_table.strip():
                    documents.append(
                        Document(
                            page_content=markdown_table,
                            metadata={
                                "source": source,
                                "page": 0,
                                "file_type": "docx",
                                "content_type": "table",
                                "table_index": table_index,
                            },
                        )
                    )

        flush_paragraphs()

        return documents

    # ------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------

    @staticmethod
    def _load_txt(file_path: str) -> list[Document]:

        source = Path(file_path).name
        loader = TextLoader(file_path, encoding="utf-8")
        documents = loader.load()

        for doc in documents:
            # TextLoader defaults metadata["source"] to the full path passed
            # in, while the PDF/DOCX loaders above use just the filename --
            # normalize here so every loader agrees on what "source" means.
            doc.metadata["source"] = source
            doc.metadata["file_type"] = "txt"
            doc.metadata["content_type"] = "text"
            doc.metadata.setdefault("page", 0)

        return documents
