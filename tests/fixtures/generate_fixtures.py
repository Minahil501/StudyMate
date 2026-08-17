"""
generate_fixtures.py

One-off script to (re)generate the 3 sample test documents used by the
test suite:

    sample.pdf   - text + a data table (Compiler Theory topic)
    sample.docx  - text + a data table (Data Engineering topic)
    sample.txt   - plain text, no table

Run with:  python tests/fixtures/generate_fixtures.py

These mirror Minahil's actual coursework topics so retrieval results in
tests/manual scripts look like real StudyMate usage instead of generic
lorem-ipsum filler.
"""

from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

FIXTURES_DIR = Path(__file__).parent


def make_pdf():
    path = FIXTURES_DIR / "sample.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("COMPILER THEORY", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Chapter 3: LL(1) Parsing", styles["Heading1"]))
    story.append(Paragraph(
        "An LL(1) parser is a top-down parser that reads input Left to right "
        "and produces a Leftmost derivation, using 1 token of lookahead. "
        "A grammar is LL(1) if, for every pair of productions A -> alpha | beta, "
        "the FIRST sets of alpha and beta are disjoint, and if epsilon is in "
        "FIRST(alpha), then FIRST(beta) and FOLLOW(A) are also disjoint.",
        styles["BodyText"],
    ))
    story.append(Spacer(1, 8))

    story.append(Paragraph("FIRST and FOLLOW Set Example", styles["Heading2"]))
    story.append(Paragraph(
        "Consider the grammar E -> T E' , E' -> + T E' | epsilon , "
        "T -> F T' , T' -> * F T' | epsilon , F -> ( E ) | id. "
        "The FIRST and FOLLOW sets for each non-terminal are shown below.",
        styles["BodyText"],
    ))
    story.append(Spacer(1, 8))

    table_data = [
        ["Non-Terminal", "FIRST Set", "FOLLOW Set"],
        ["E", "{ ( , id }", "{ ) , $ }"],
        ["E'", "{ + , epsilon }", "{ ) , $ }"],
        ["T", "{ ( , id }", "{ + , ) , $ }"],
        ["T'", "{ * , epsilon }", "{ + , ) , $ }"],
        ["F", "{ ( , id }", "{ * , + , ) , $ }"],
    ]

    table = Table(table_data, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
    ]))

    story.append(table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Chapter 4: Panic Mode Error Recovery", styles["Heading1"]))
    story.append(Paragraph(
        "Panic mode recovery discards input symbols one at a time until a "
        "token in a predefined synchronizing set is found. Synchronizing "
        "tokens are usually statement or block delimiters such as semicolons "
        "or closing braces, chosen so the parser can resume without cascading "
        "further errors.",
        styles["BodyText"],
    ))

    doc.build(story)
    print(f"wrote {path}")


def make_docx():
    path = FIXTURES_DIR / "sample.docx"
    doc = DocxDocument()

    doc.add_heading("Data Engineering / Big Data", level=0)

    doc.add_heading("Topic: The CAP Theorem", level=1)
    doc.add_paragraph(
        "The CAP theorem states that a distributed data store can provide "
        "at most two of the following three guarantees simultaneously: "
        "Consistency, Availability, and Partition tolerance. Since network "
        "partitions are unavoidable in real distributed systems, practical "
        "designs choose between CP and AP behavior during a partition."
    )

    doc.add_heading("Comparison of Distributed Database Models", level=2)
    doc.add_paragraph(
        "The table below compares common database systems by which two "
        "CAP guarantees they prioritize."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Database"
    header[1].text = "CAP Priority"
    header[2].text = "Typical Use Case"

    rows = [
        ("MongoDB", "CP", "Document storage with strong consistency"),
        ("Cassandra", "AP", "High write-throughput, multi-region"),
        ("HBase", "CP", "Hadoop-based wide-column storage"),
        ("DynamoDB", "AP", "Highly available key-value store"),
    ]

    for db, priority, use_case in rows:
        cells = table.add_row().cells
        cells[0].text = db
        cells[1].text = priority
        cells[2].text = use_case

    doc.add_heading("Topic: Batch vs Stream Processing", level=1)
    doc.add_paragraph(
        "Batch processing operates on bounded datasets collected over a "
        "period of time (e.g. nightly ETL jobs), while stream processing "
        "operates on unbounded data arriving continuously, producing "
        "results with much lower latency at the cost of more complex "
        "state management."
    )

    doc.save(path)
    print(f"wrote {path}")


def make_txt():
    path = FIXTURES_DIR / "sample.txt"
    content = """IEEE CS UCP Student Branch - Think2Code 2.0 Event Notes

Overview
--------
Think2Code 2.0 is a competitive programming event organized to give
students hands-on experience with algorithmic problem solving under
time pressure, similar in spirit to contests like Codeforces or ICPC
regionals.

Rules
-----
- Teams of up to 2 members.
- 5 problems, 3 hours.
- Partial scoring is enabled for partially correct submissions.
- Standard time/memory limits: 1-2 seconds, 256 MB.

Judging Criteria
-----------------
Submissions are judged automatically against hidden test cases. Ties
on score are broken by total submission time, with a penalty added
for every incorrect attempt on a problem that is eventually solved.

Why This Matters for Preparation
---------------------------------
Practicing under a similar format (multiple problems, a hard time
limit, and a penalty for wrong attempts) is the best way to simulate
contest pressure before the actual event.
"""
    path.write_text(content, encoding="utf-8")
    print(f"wrote {path}")


if __name__ == "__main__":
    make_pdf()
    make_docx()
    make_txt()
